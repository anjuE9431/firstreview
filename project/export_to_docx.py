import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

project_root = "."  # current folder
output_file = "fish_market_app_code.docx"

doc = Document()
doc.add_heading("Fish Market App - Full Source Code", level=1)

for root, dirs, files in os.walk(project_root):
    for file in files:
        if file.endswith((".py", ".html", ".css", ".js", ".sql", ".txt")):
            filepath = os.path.join(root, file)
            relpath = os.path.relpath(filepath, project_root)

            doc.add_heading(relpath, level=2)
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                code = f.read()

            paragraph = doc.add_paragraph()
            run = paragraph.add_run(code)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            # Make sure font applies in Word (especially for non-Latin systems)
            r = run._element.rPr.rFonts
            r.set(qn('w:eastAsia'), 'Courier New')

doc.save(output_file)
print("✅ Export complete:", output_file)
